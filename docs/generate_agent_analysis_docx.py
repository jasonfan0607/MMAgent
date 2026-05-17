r"""生成 MathModelAgent Agent 实现分析的 Word 文档。

使用方法（项目用 uv 管理依赖）：
    cd D:\MathModelAgent\backend
    uv pip install python-docx
    .\.venv\Scripts\python.exe ..\docs\generate_agent_analysis_docx.py

输出：docs/MathModelAgent_Agent实现分析.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ---- 样式辅助 ----


def _set_default_font(doc: Document, font_name: str = "微软雅黑") -> None:
    """把 Normal 样式设为中文友好字体。"""
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(11)
    # 让中文走指定字体
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Consolas")
    r_fonts.set(qn("w:hAnsi"), "Consolas")


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    if level == 0:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = "微软雅黑"
        rpr = run._element.get_or_add_rPr()
        r_fonts = rpr.find(qn("w:rFonts"))
        if r_fonts is None:
            from docx.oxml import OxmlElement

            r_fonts = OxmlElement("w:rFonts")
            rpr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "微软雅黑")


def add_paragraph(doc: Document, text: str) -> None:
    """支持简单内联 markdown：**bold** 和 `code`。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _append_inline(p, text)


def _append_inline(p, text: str) -> None:
    import re

    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            run = p.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = p.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    _append_inline(p, text)


def add_code_block(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1B, 0x1F, 0x23)

    # 浅灰底
    from docx.oxml import OxmlElement

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if header and i == 0:
                run.bold = True


# ---- 文档内容 ----


def build_document(output_path: Path) -> None:
    doc = Document()
    _set_default_font(doc)

    add_heading(doc, "MathModelAgent 项目 Agent 实现深度分析", 0)

    intro = (
        "本文档系统拆解 MathModelAgent 项目中 Agent 的实现方式，"
        "覆盖整体架构、Agent 基类、四个具体 Agent、LLM 抽象层、工具系统、"
        "Agent 间通信和工作流编排七个层面。"
    )
    add_paragraph(doc, intro)

    # ---- 1 ----
    add_heading(doc, "一、整体架构：四 Agent 流水线", 1)
    add_paragraph(
        doc,
        "入口是 `MathModelWorkFlow`（`backend/app/core/workflow.py`），"
        "整个建模任务是一条严格的流水线：",
    )
    pipeline = (
        "用户题目\n"
        "   ↓\n"
        "CoordinatorAgent  → 识别意图、拆解问题为结构化 JSON\n"
        "   ↓\n"
        "ModelerAgent      → 为每个子问题给出建模方案（不写代码）\n"
        "   ↓\n"
        "CoderAgent  ──工具调用→ 本地 Jupyter / E2B 沙箱\n"
        "   ↓ (循环求解每个 ques)\n"
        "WriterAgent ──工具调用→ OpenAlex 文献搜索\n"
        "   ↓ (按论文章节循环)\n"
        "最终论文"
    )
    add_code_block(doc, pipeline, caption="流水线示意")
    add_paragraph(
        doc,
        "每个 Agent 持有自己**独立的 `LLM` 实例**（可以是不同的模型/服务商），"
        "由 `LLMFactory` 一次性创建。这种「一个 Agent 一个 LLM 配置」的设计让用户"
        "可以给协调者用便宜模型、给建模手/写作手用强推理模型、给代码手用擅长 Tool Use 的模型。",
    )

    # ---- 2 ----
    add_heading(doc, "二、Agent 基类：对话历史 + Token 感知的记忆压缩", 1)
    add_paragraph(
        doc,
        "所有 Agent 继承 `app.core.agents.agent.Agent`，"
        "基类只做一件事：**管理对话历史并在快撑爆上下文窗口前自动压缩**。",
    )

    add_heading(doc, "2.1 核心字段", 2)
    add_code_block(
        doc,
        '''class Agent:
    """Agent 基类，管理对话历史、轮次控制和记忆压缩。"""

    def __init__(
        self,
        task_id: str,
        model: LLM,
        context_window: int = 128000,  # 模型上下文窗口大小（token）
        token_threshold_ratio: float = _DEFAULT_TOKEN_THRESHOLD_RATIO,
    ) -> None:
        self.task_id = task_id
        self.model = model
        self.chat_history: list[dict] = []
        self.context_window = context_window
        self.token_threshold_ratio = token_threshold_ratio
        self.current_token_count = 0''',
        caption="backend/app/core/agents/agent.py L15-L30",
    )
    add_bullet(
        doc,
        "`chat_history` 用 OpenAI 风格的消息字典数组保存"
        "（`role` + `content` + 可选 `tool_calls`/`tool_call_id`/`reasoning_content`）。",
    )
    add_bullet(
        doc,
        "`current_token_count` 优先使用 API 返回的真实 `prompt_tokens`，"
        "没有时才用 `len(text) // 3` 的本地估算。",
    )
    add_bullet(doc, "`context_window` 默认 128k，触发压缩的阈值是它的 75%。")

    add_heading(doc, "2.2 append + 自动压缩", 2)
    add_paragraph(
        doc,
        "每次往历史里加消息都走 `append_chat_history`，加完后立即检查阈值。"
        "**只有 `role != \"tool\"` 时才会触发压缩**，这是关键工程细节——"
        "绝不能在 `assistant.tool_calls` 和对应的 `tool` 响应之间把历史切断，"
        "否则下一轮 API 调用会直接报「tool_call_id not found」。",
    )
    add_code_block(
        doc,
        '''async def append_chat_history(self, msg: dict) -> None:
    self.chat_history.append(msg)
    self.current_token_count += self._estimate_message_tokens(msg)
    # 只有在添加非tool消息时才进行内存清理，避免在工具调用期间破坏消息结构
    if msg.get("role") != "tool":
        await self.compress_if_needed()''',
    )

    add_heading(doc, "2.3 压缩策略", 2)
    add_paragraph(doc, "`compress_if_needed` 的算法是：")
    for line in [
        "保留首条 `system` 消息；",
        "用 `_find_safe_preserve_point` 从尾部往前找一个**不会切断工具调用对**的位置，保留最后若干条；",
        "把「系统消息和保留区之间」的所有历史扔给 LLM 用 `simple_chat` 做总结；",
        "重组为 `[system, \"[历史对话总结] …\", …保留消息]`；",
        "如果总结调用失败，回退到 `_get_safe_fallback_history()`，按完整工具调用对截取。",
    ]:
        add_bullet(doc, line)
    add_paragraph(
        doc,
        "`_is_safe_cut_point` 的判断是：**从切点开始向后扫描，每个 `role==\"tool\"` "
        "的消息都必须能在切点之后找到对应的 `tool_calls.id`**，否则该切点不安全。",
    )

    # ---- 3 ----
    add_heading(doc, "三、四个具体 Agent", 1)

    add_heading(doc, "3.1 CoordinatorAgent — 解析意图，输出结构化 JSON", 2)
    add_paragraph(doc, "工作模式：")
    for line in [
        "注入 `COORDINATOR_PROMPT`（要求输出固定 schema：`title / background / ques_count / ques1..quesN`）；",
        "**无限循环调用 LLM 直到 JSON 解析成功**——失败时把错误描述拼回 system prompt 重试；",
        "返回 Pydantic 模型 `CoordinatorToModeler`。",
    ]:
        add_bullet(doc, line)
    add_code_block(
        doc,
        '''except (json.JSONDecodeError, ValueError, KeyError) as e:
    attempt += 1
    error_prompt = f"⚠️ 上次响应格式错误: {str(e)}。请严格输出JSON格式"
    await self.append_chat_history({
        "role": "system",
        "content": self.system_prompt + "\\n" + error_prompt
    })''',
    )
    add_paragraph(doc, "它**不使用工具调用**，纯粹是 schema 提取器。")

    add_heading(doc, "3.2 ModelerAgent — 出方案，不出代码", 2)
    add_paragraph(
        doc,
        "签名 `run(coordinator_to_modeler) -> ModelerToCoder`。"
        "它额外包了一个 `repair_json` 函数做三级降级修复：直接 parse → "
        "正则修字符串里的未转义引号 → 用键值对正则兜底提取。",
    )
    add_paragraph(
        doc,
        "`MODELER_PROMPT` 强制输出 schema "
        "`{\"eda\": ..., \"ques1\": ..., \"quesN\": ..., \"sensitivity_analysis\": ...}`，"
        "每个 value 是一段建模方案文字，**全程禁止给代码**——这是和 CoderAgent 的分工边界。",
    )

    add_heading(doc, "3.3 CoderAgent — 核心 Tool-Use 智能体", 2)
    add_paragraph(
        doc,
        "`CoderAgent.run` 是整个项目最复杂的 Agent 循环，"
        "结构是经典的 **ReAct + 反思 + 双重上限**。它同时跟踪两个计数：",
    )
    add_bullet(doc, "`current_chat_turns`：累计对话轮数（防止单个子任务死循环）；")
    add_bullet(doc, "`retry_count`：连续报错次数（决定是否再反思）。")
    add_paragraph(doc, "**首次运行时**注入 `CODER_PROMPT` 和当前工作目录下的文件列表：")
    add_code_block(
        doc,
        '''if self.is_first_run:
    self.is_first_run = False
    await self.append_chat_history(
        {"role": "system", "content": self.system_prompt}
    )
    await self.append_chat_history(
        {
            "role": "user",
            "content": f"当前文件夹下的数据集文件{get_current_files(self.work_dir, 'data')}",
        }
    )''',
    )
    add_paragraph(
        doc,
        "注意 `is_first_run` 这个标志位——"
        "**CoderAgent 在跨子任务时复用同一个实例和同一份 `chat_history`**，"
        "这样代码手能记住前一题用了什么变量、跑过哪些代码（Jupyter 内核也是同一个，变量保留在内存里）。",
    )

    add_paragraph(doc, "单轮循环要点：")
    for line in [
        "**退出条件不是 prompt，而是「LLM 在某一轮没有再发起 `tool_calls`」**。模型自己决定什么时候停。",
        "出错时不是直接抛——`get_reflection_prompt(error_message, code)` 会生成一段「分析错误→修代码→重试」的反思提示注入 history。",
        "成功执行后什么都不加，直接 `continue` 让 LLM 决定「再写一段」还是「结束」。",
        "历史里的工具调用对结构严格遵循 OpenAI 规范：`assistant{tool_calls} → tool{tool_call_id, content}`。",
    ]:
        add_bullet(doc, line)
    add_paragraph(
        doc,
        "工具定义只有一个 `execute_code`。"
        "代码手返回值通过 `CoderToWriter` 把「最终的文字回答 + 这一节产出的图片文件名列表」"
        "打包传给写作手。图片列表是通过解释器内部用 `last_created_images` 集合做差集算出来的。",
    )

    add_heading(doc, "3.4 WriterAgent — 写作 + 文献检索", 2)
    add_paragraph(
        doc,
        "WriterAgent 也是 Tool Use 模式，但只有 **search_papers** 一个工具，"
        "而且**不循环**（只处理至多一次工具调用）。设计取舍很清楚："
        "**论文写作不是 ReAct，而是「先查文献→再下笔」的两段式**。",
    )
    add_paragraph(
        doc,
        "`available_images` 列表会被拼成一段 markdown 强制要求 LLM 把图插进去，"
        "避免论文里只写「如图所示」但没有 `![](xxx.png)`。",
    )

    # ---- 4 ----
    add_heading(doc, "四、LLM 抽象层：Provider 模式", 1)
    add_paragraph(
        doc,
        "四个 Agent 都通过 `LLM` 类访问模型。`LLM` 自己**不发请求**，它持有一个 `BaseProvider` 实例：",
    )
    add_code_block(
        doc,
        '''def _create_provider(self, api_type: ApiType | None) -> BaseProvider:
    match api_type:
        case ApiType.OPENAI_RESPONSES:
            return OpenAIResponsesProvider()
        case ApiType.ANTHROPIC:
            return AnthropicProvider()
        case _:
            return OpenAIChatProvider()''',
    )
    add_paragraph(
        doc,
        "三个 Provider 都实现 `BaseProvider.call(...) -> StandardResponse`，"
        "把不同 API 的返回值归一化成同一个数据类 `StandardResponse`，"
        "包含 `content / reasoning_content / tool_calls / usage` 四个字段。"
        "这样 Agent 代码里写 `response.tool_calls[0].arguments` 永远是合法的，"
        "不管底层是 OpenAI 还是 Anthropic。",
    )

    add_heading(doc, "4.1 LLM.chat 的附加职责", 2)
    add_paragraph(doc, "`LLM.chat` 除了调 provider 还做三件事：")
    add_bullet(
        doc,
        "**`_validate_and_fix_tool_calls`**：调用前再扫一遍历史，删掉没有对应 "
        "`tool` 响应的悬空 `tool_calls`，以及找不到归宿的孤儿 `tool` 消息。",
    )
    add_bullet(
        doc,
        "**重试**：抓异常 → `time.sleep(retry_delay * min(attempt, 10))` → "
        "重试，默认无最大次数。",
    )
    add_bullet(
        doc,
        "**Redis 广播**：把每一次 LLM 输出按 agent 类型包成 "
        "`CoderMessage / WriterMessage / ModelerMessage / CoordinatorMessage` "
        "推到 Redis，前端 WebSocket 实时显示。",
    )

    # ---- 5 ----
    add_heading(doc, "五、Agent 间通信：Pydantic Schema 而不是消息总线", 1)
    add_paragraph(
        doc,
        "虽然项目里有 Redis pubsub，但**那是给前端看的**。"
        "Agent 之间的「交接班」用的是 `app/schemas/A2A.py` 里的 Pydantic 模型：",
    )
    add_code_block(
        doc,
        '''class CoordinatorToModeler(BaseModel):
    questions: dict
    ques_count: int

class ModelerToCoder(BaseModel):
    questions_solution: dict[str, str]

class CoderToWriter(BaseModel):
    code_response: str | None = None
    code_output: str | None = None
    created_images: list[str] | None = None

class WriterResponse(BaseModel):
    response_content: Any
    footnotes: list[tuple[str, str]] | None = None''',
    )
    add_paragraph(
        doc,
        "所以本质上**这是一个用类型来约束的顺序管道**，不是事件驱动也不是黑板模型。"
        "文件名「A2A」是 Agent-to-Agent 的缩写，但实际上是同步调用 + 强类型 DTO。",
    )

    # ---- 6 ----
    add_heading(doc, "六、工作流编排：Flows 把 prompt 组装外提", 1)
    add_paragraph(
        doc,
        "`MathModelWorkFlow.execute` 里没有任何 prompt 字符串拼接，"
        "所有「给 CoderAgent 的子任务 prompt」和「给 WriterAgent 的章节 prompt」"
        "都通过 `Flows` 类生成。`Flows` 把整个写作任务的顺序固化为：",
    )
    add_code_block(
        doc,
        '''seq = [
    "firstPage",
    "RepeatQues",
    "analysisQues",
    "modelAssumption",
    "symbol",
    "eda",
    *ques_str,
    "sensitivity_analysis",
    "judge",
]''',
    )
    add_paragraph(
        doc,
        "求解阶段只跑 `eda → ques1…quesN → sensitivity_analysis`（数据相关），"
        "写作阶段才补齐封面、问题重述、模型评价等不需要代码的章节。",
    )

    # ---- 7 ----
    add_heading(doc, "七、Tool 层：代码解释器是「Agent 的手」", 1)
    add_paragraph(
        doc,
        "CoderAgent 的 `execute_code` 工具背后是 `BaseCodeInterpreter` 抽象，"
        "由工厂决定用本地 Jupyter 还是 E2B 沙箱。"
        "`LocalCodeInterpreter` 用 `jupyter_client` 启一个长期内核，"
        "预先注入中文字体加载和工作目录切换代码。",
    )
    add_paragraph(
        doc,
        "`execute_code` 返回三元组 `(text_to_gpt, error_occurred, error_message)`，"
        "这正是 CoderAgent 循环里 `if error_occurred: ...` 判断的来源。",
    )
    add_paragraph(doc, "每次执行的输出会三路分发：")
    add_bullet(doc, "拼成纯文本返回给 LLM（图片用 `[image]` 占位，因为 LLM 看不到 base64）；")
    add_bullet(doc, "写进 `notebook_serializer` 生成 `.ipynb` 落盘；")
    add_bullet(doc, "通过 `_push_to_websocket` 推 Redis 给前端实时显示。")
    add_paragraph(
        doc,
        "`add_section` / `get_code_output` / `get_created_images` 这套接口是按「章节」组织的——"
        "和 Flows 的 key 一一对应（`eda` / `ques1` / `sensitivity_analysis`），"
        "所以 WriterAgent 写哪一章就能取出哪一章的代码输出和图片。",
    )

    # ---- 8 ----
    add_heading(doc, "八、值得关注的工程取舍", 1)
    rows = [
        ["维度", "做法", "意图"],
        [
            "Agent 生命周期",
            "跨子任务复用同一实例 + Jupyter 内核共享",
            "代码手能记住上一题的变量和结论",
        ],
        [
            "框架依赖",
            "不依赖 LangChain/AutoGen/CrewAI，手写 while-True",
            "降低魔法，调试容易",
        ],
        [
            "结构化输出",
            "Prompt 约束 + 正则清洗 + 错误反馈重试 + repair_json 三级降级",
            "不能假设 LLM 一次输出合法 JSON",
        ],
        [
            "Token 估算",
            "真实 prompt_tokens 优先，回退到 `len/3` 估算",
            "兼顾准确性和容错",
        ],
        [
            "记忆压缩",
            "调用 LLM 做摘要 + 保护工具调用对完整性",
            "永远不会切坏 assistant.tool_calls 与 tool 响应的配对",
        ],
        [
            "领域知识",
            "把竞赛经验固化进 CODER_PROMPT / MODELER_PROMPT",
            "把项目「软资产」沉淀到 prompt 而非依赖模型自带知识",
        ],
    ]
    add_table(doc, rows, header=True)

    # ---- 9 ----
    add_heading(doc, "九、一句话总结", 1)
    add_paragraph(
        doc,
        "MathModelAgent 实现 Agent 的方式是：**继承一个负责对话历史 + token 感知压缩的基类 → "
        "每个 Agent 在 `run()` 里手写自己的状态机（CoordinatorAgent 是「重试直到 JSON 合法」、"
        "ModelerAgent 类似、CoderAgent 是「ReAct + 反思 + 双上限」、"
        "WriterAgent 是「一次性工具调用 + 文献检索」）→ "
        "LLM 调用通过 Provider 抽象层统一成 `StandardResponse` → "
        "Agent 之间用 Pydantic DTO 同步交接 → "
        "一个 `MathModelWorkFlow` 按固定顺序串起来，配合 Flows 把 prompt 组装逻辑外提**。",
    )
    add_paragraph(
        doc,
        "整体没有用任何 Agent 框架，全部手写，"
        "所以可以非常清晰地看到每一条消息、每一次工具调用、每一次记忆压缩到底在做什么。",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "MathModelAgent_Agent实现分析.docx"
    build_document(out)
